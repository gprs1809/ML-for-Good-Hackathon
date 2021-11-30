#!/usr/bin/env python
# coding: utf-8


import pandas as pd
import os
import textract
import re
import gensim
from gensim.summarization.summarizer import summarize
from gensim.summarization import keywords
from rake_nltk import Rake
import nltk
nltk.download('stopwords')
nltk.download('punkt')
from sentence_transformers import SentenceTransformer
#sbert_model = SentenceTransformer('bert-base-nli-mean-tokens')
import tensorflow as tf
import tensorflow_hub
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import RandomForestClassifier
from summarizer import Summarizer
import docx
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer



model = Summarizer()
uni_encoder = tensorflow_hub.KerasLayer('https://tfhub.dev/google/universal-sentence-encoder/4',trainable=False)



#function for preprocessing of word documents comprising of dialogues between Moderators, Parents and Administrator
def doc_preprocessing(path):
    """ 
    Code for preprocessing docx files under focusGroups. The steps involved are:
    1) Remove irrelevant strings like (silence, [inaudible], next line (\n), extra tabs (\t) etc.)
    2) We collect each dialogue from the speakers: Parents, Moderators, Speaker, Adminitrator.
    3) We ultimately want only the Parents'dialogues to get more insights on the impact of lockdown and tech on their children.
    4) All the dialogues of a Parent are combined for further analysis.
    5) The function returns a dictionary with keys as Parents and the values as their combined dialogue from the entire doc.
    """
    text = textract.process(path)
    text = text.decode('utf-8')
    text =text.replace('\n',"")
    text = text.replace('\t',"")
    text = text.replace('(silence)',"")
    text = text.replace('. ','.')
    split_sent=re.split(r'[?.]', text)
    split_sent = list(filter(('').__ne__, split_sent))
    split_sent = [sent for sent in split_sent if not re.search("inaudible",sent)]
    dial_characters = ('Administrator','Moderator','Parent','Speaker',' Administrator',' Moderator',' Parent',' Speaker')
    dial_start=[ind for ind,l in enumerate(split_sent) if l.startswith(dial_characters)]
    list_of_dial = []
    for i in range(len(dial_start)-1):
        list_of_dial.append('.'.join(split_sent[dial_start[i]:dial_start[i+1]]))
    list_of_Parent_dial = [sent for sent in list_of_dial if sent.startswith('Parent')]
    num_ind = [sent.split(" ")[1].split(':')[0] for sent in list_of_Parent_dial]
    parent_num = list(set([int(word) for sent in num_ind for word in sent.split() if word.isdigit()]))
    all_dial_comb = []
    for i in parent_num:
        all_dial_comb.append('.'.join([sent for sent in list_of_Parent_dial if sent.startswith(('Parent '+str(i)+':',' Parent '+str(i)+':'))]))
    all_d_comb_n=[]
    for i in range(len(all_dial_comb)):
        all_d_comb_n.append(all_dial_comb[i].replace('Parent '+str(i+1)+":","").replace(' Parent '+str(i+1)+":",""))
    dial_dict = dict(zip(['Parent '+str(i) for i in parent_num],all_d_comb_n))
    for key in list(dial_dict.keys()):
        dial_dict[key] = dial_dict.get(key).replace("My name's Parent",'').replace('Parent','').replace('Moderator','').replace('Administrator','')
    #dial_dict[key] = re.sub('[0-9]','',dial_dict.get(key))
    return dial_dict





# we can draw insights from the extracted summaries of dialogues. Either this could be our result or we can work on these summaries further
#for say, sentiment analysis or clustering.
def extractive_text_summarization(path):
    """
    This function uses the processed data from doc_preprocessing function. We perform an extractive text summarization for each Parents dialogues.
    We have used the summarizer from gensim package and we also use the Bert summarizer. With these packages, we can extract important text based on word 
    count, percentage, number of sentences. We had tried key phrase extraction which didn't work well.
    """
    imp_dial_parents = doc_preprocessing(path)
    # Summary (0.5% of the original content).
    summ_per = [summarize(dialogue.replace(".",". "), ratio = 0.1) for dialogue in list(imp_dial_parents.values())]
    # Summary (200 words)
    summ_words = [summarize(dialogue.replace(".",". "), word_count = 200) for dialogue in list(imp_dial_parents.values())]
    keys = list(imp_dial_parents.keys())
    summ_per = dict(zip(keys,summ_per))
    summ_words = dict(zip(keys,summ_words))
    result1 = [model(dialogue, ratio=0.2) for dialogue in list(imp_dial_parents.values())]   # Specified with ratio
    result2 = [model(dialogue, num_sentences=6) for dialogue in list(imp_dial_parents.values())]
    bert_ratio = dict(zip(keys,result1))
    bert_num_sent = dict(zip(keys,result2))
    # r = Rake()
    # key_ph_list = []
    # for i in range(len(keys)):
    #   r.extract_keywords_from_text(list(imp_dial_parents.values())[i])
    #   key_ph_list.append(r.get_ranked_phrases()[:30])
    # key_phrases = dict(zip(keys,key_ph_list))
    return summ_per, summ_words, bert_ratio, bert_num_sent





list_of_docs = ['./ML-for-Good-Hackathon/Data/FocusGroups/'+file for file in os.listdir('./ML-for-Good-Hackathon/Data/FocusGroups')]



#Executing the extractive_text_summarization function and saving the outputs
para_topics = ['gensim_summ_ratio', 'gensim_summ_words', 'bert_summ_ratio', 'bert_summ_num_sent']
for doc in list_of_docs:
    extract = extractive_text_summarization(doc)
    mydoc = docx.Document()
    for i in range(len(extract)):
        mydoc.add_paragraph(para_topics[i]+':')
        for key in extract[i].keys():
            mydoc.add_paragraph(key+': '+extract[i].get(key))
    mydoc.save('summ_outputs_'+doc.split('/')[-1])



def preprocess_prolificacademic_and_feature_importances(path):
  """
  only works for Nov 2020 and April 2021 data since the data before Nov 2020 doesn't contain the variable suspectedinfected. We set the target
  feature as suspectedidentified and fit a random forest classifier on the processed dataset. Note that the purpose of the classifier is not to
  predict because we realize that the features in the data need not be same always. The purpose is to just suggest the more important features in
  the processed dataset that can. The preprocessing steps are the following:
  1) remove features with more than 50% NaNs.
  2) find the top 19 or 20 features that are most correlated to suspectedidentified among each of int and float features.
  3) replace NaNs in the string features with 'Information not available'.
  4) combine dataframes from 2) and 3).
  5) remove all rows with NaNs.
  6) set target as suspectedidentified. 
  7) for the string features, get embeddings for each using the Universal sentence encoder pretrained model from tensorflow hub.
  8) For each of the string features, step 7) would give us embeddings of dimension 571. We perform PCA on each to reduce the dimension to 5 for each.
  9) In 8), the new features as embeddings are given names as feature_name followed by an index.
  10) By step 9), we are done with preprocessing. We then fit a random forest classifier to arrive at the important features. This might help Child Mind
  Institute to focus on a fewer features. In case of embeddings, the features for which some components of the embeddings are important can be focussed on.

"""
    data2 = pd.read_csv(path)
    to_be_removed = list((data2.isnull().sum() * 100 / len(data2) >50)[data2.isnull().sum() * 100 / len(data2) >50].index)
    time_col = [col for col in data2.columns if col.startswith('timestamp')]
    to_be_removed = to_be_removed+['ID']+time_col
    data2.drop(to_be_removed, axis=1, inplace=True)
    g = data2.columns.to_series().groupby(data2.dtypes).groups
    g_dic = {k.name: v for k, v in g.items()}
    float_f = list(g_dic.get('float64'))
    f_var = data2[float_f]
    f_var_cor = f_var.corr()
    f_var_list = list(f_var_cor.suspectedinfected.sort_values(ascending=False)[1:20].index)
    print(f_var_list)
    f_var = f_var[f_var_list]
    int_f = list(g_dic.get('int64'))
    int_var = data2[int_f+['suspectedinfected']]
    i_var_corr = int_var.corr()
    i_var_list = list(i_var_corr.suspectedinfected.sort_values(ascending=False)[0:20].index)
    print(i_var_list)
    int_var = int_var[i_var_list]
    #int_var.drop(['suspectedinfected'], axis=1,inplace=True)
    var_to_preprocess = list(g_dic.get('object'))
    string_f = data2[var_to_preprocess]
    string_f=string_f.fillna('information not available')
    combined_df = pd.concat([f_var, int_var,string_f], axis=1)
    print(f"'% of rows usable for training':{len(combined_df.dropna())/len(combined_df)}")
    combined_df.dropna(inplace=True)
    target = combined_df.pop('suspectedinfected')
    print(len(combined_df))
    for feat in var_to_preprocess:
        a=uni_encoder(list(combined_df[feat]))
        principal=PCA(n_components=5)
        principal.fit(a)
        x=principal.transform(a)
        y = pd.DataFrame(x,columns=[feat+str(i) for i in range(1,6)],index=combined_df.index)
        combined_df = pd.concat([combined_df,y],axis=1)
    combined_df.drop(var_to_preprocess,axis=1,inplace=True)
    clf = RandomForestClassifier(n_estimators=500,max_depth=10,random_state=42)
    clf.fit(combined_df,target)
    feats = {} # a dict to hold feature_name: feature_importance
    for feature, importance in zip(combined_df.columns, clf.feature_importances_):
        feats[feature] = importance #add the name/value pair 

    importances = pd.DataFrame.from_dict(feats, orient='index').rename(columns={0: 'Gini-importance'})
    importances = importances.sort_values(by='Gini-importance',ascending=False)
    #importances.head(50)
    return importances.head(50)




#Executing preprocess_prolificacademic_and_feature_importances for Nov 2020 and April 2021 data (we have done it for both
#the previous as well as updated data since we had already done it with the previous data)
April2021_adult_feat_imp=preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/April 2021/Data/CRISIS_Adult_April_2021.csv')
Nov2020_adult_feat_imp=preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/November 2020/Data/CRISIS_Adult_November_2020.csv')
April2021_parent_feat_imp=preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/April 2021/Data/CRISIS_Parent_April_2021.csv')
Nov2020_parent_feat_imp=preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/November 2020/Data/CRISIS_Parent_November_2020.csv')
updated_April2021_adult_feat_imp=preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/updated_data/April_21/CRISIS_Adult_April_2021.csv')
updated_April2021_parent_feat_imp=preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/updated_data/April_21/CRISIS_Parent_April_2021.csv')
updated_Nov2020_adult_feat_imp = preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/updated_data/November/CRISIS_Adult_November_2020.csv')
updated_Nov2020_parent_feat_imp = preprocess_prolificacademic_and_feature_importances('./ML-for-Good-Hackathon/Data/ProlificAcademic/updated_data/November/CRISIS_Parent_November_2020.csv')




#saving results from preprocess_prolificacademic_and_feature_importances in a single xlsx file
def save_xls(list_dfs, xls_path,required_dfs):
    with pd.ExcelWriter(xls_path) as writer:
        for n, df in enumerate(list_dfs):
            df.to_excel(writer,required_dfs[n])
        writer.save()

#list_dfs = [April2021_adult_feat_imp, April2021_parent_feat_imp, Nov2020_adult_feat_imp,Nov2020_parent_feat_imp,updated_April2021_adult_feat_imp,updated_April2021_parent_feat_imp,updated_Nov2020_adult_feat_imp,updated_Nov2020_parent_feat_imp]
xls_path = '/content/prolificacademic_feature_imp.xlsx'

required_dfs = [df for df in globals() if 'feat_imp' in df]

#creating the list of dataframes in the same order as their names as string are listed in required_dfs
list_dfs = [April2021_adult_feat_imp,
 Nov2020_adult_feat_imp,
 April2021_parent_feat_imp,
 Nov2020_parent_feat_imp,
 updated_April2021_adult_feat_imp,
 updated_April2021_parent_feat_imp,
 updated_Nov2020_adult_feat_imp,
 updated_Nov2020_parent_feat_imp]

required_dfs = [df[:31] if len(df) > 31  else df for df in required_dfs]

save_xls(list_dfs,xls_path,required_dfs)




def sentiment_scores(sentence):
    """
    getting sentiment scores for every sentence in crisislogger file.
    """
    sid_obj = SentimentIntensityAnalyzer()
    sentiment_dict=sid_obj.polarity_scores(sentence)
    Negative_score = sentiment_dict['neg']*100
    Neutral_score = sentiment_dict['neu']*100
    Positive_score = sentiment_dict['pos']*100
    if (sentiment_dict['compound'] >= 0.05):
        Overall_rate = 'Positive'
    elif (sentiment_dict['compound'] <= - 0.05):
        Overall_rate = 'Negative'
    else :
        Overall_rate = 'Neutral'
    return Negative_score,Neutral_score,Positive_score,Overall_rate



#executing sentiment_scores on crisislogger data
data = pd.read_csv(r"./ML-for-Good-Hackathon/Data/CrisisLogger/crisislogger.csv") 
res = data['transcriptions'].apply(sentiment_scores)
data['Negative_score'] = [x[0] for x in res]
data['Neutral_score'] = [x[1] for x in res]
data['Positive_score'] = [x[2] for x in res]
data['Overall_rate']=[x[3] for x in res]
data.to_excel('./sentiment_scores_crisislogger.xlsx')

